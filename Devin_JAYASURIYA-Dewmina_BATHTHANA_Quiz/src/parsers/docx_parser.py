"""Parseur pour les documents Word (.docx)."""

from pathlib import Path
from typing import List, Dict, Any
from docx import Document

from .base_parser import BaseParser


class DocxParser(BaseParser):
    """Parseur pour les documents Word."""

    def parse(self) -> List[Dict[str, Any]]:
        """Parse le document Word et retourne une liste de sections."""
        if not self.validate():
            raise FileNotFoundError(f"Le fichier {self.file_path} n'existe pas.")

        document = Document(self.file_path)
        sections = []

        for para in document.paragraphs:
            if para.style.name.startswith('Heading'):
                # Nouvelle section
                section = {
                    "title": para.text,
                    "level": int(para.style.name[-1]) if para.style.name[-1].isdigit() else 1,
                    "content": "",
                    "text": para.text
                }
                sections.append(section)
            else:
                # Contenu de section
                if sections:
                    sections[-1]["content"] += para.text + "\n"
                    sections[-1]["text"] += para.text + "\n"
                else:
                    # Contenu avant toute section
                    sections.append({
                        "title": "Introduction",
                        "level": 1,
                        "content": para.text + "\n",
                        "text": para.text + "\n"
                    })

        return sections

    def extract_text(self) -> str:
        """Extract tout le texte du document."""
        if not self.validate():
            raise FileNotFoundError(f"Le fichier {self.file_path} n'existe pas.")

        document = Document(self.file_path)
        text = []

        for para in document.paragraphs:
            text.append(para.text)

        return "\n".join(text)
